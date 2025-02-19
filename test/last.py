import sys
from PyQt5 import QtWidgets, uic
from docx import Document  
from docx.shared import RGBColor, Pt  

Form, _ = uic.loadUiType("last.ui")

class LastApp(QtWidgets.QMainWindow, Form):
    def __init__(self, score, time_taken, right_count, wrong_count):
        super().__init__()
        self.setupUi(self)

        
        self.lbl.hide()
        self.lbl1.hide()
        self.right.hide()
        self.bhim.hide()

        
        self.score = score
        self.time_taken = time_taken
        self.right_count = right_count
        self.wrong_count = wrong_count

        
        self.box1.currentIndexChanged.connect(self.check_fields)
        self.name.textChanged.connect(self.check_fields)
        self.surname.textChanged.connect(self.check_fields)
        self.section.textChanged.connect(self.check_fields)
        self.code.textChanged.connect(self.check_fields)  
        self.day.currentIndexChanged.connect(self.check_fields)  
        self.month.currentIndexChanged.connect(self.check_fields)  
        self.year.currentIndexChanged.connect(self.check_fields) 

        
        self.save.clicked.connect(self.save_to_docx)

    def check_fields(self):
        """ Show results only when the user fills all fields and selects valid options. """
        print(f"box1 index: {self.box1.currentIndex()}")
        print(f"day index: {self.day.currentIndex()}")
        print(f"month index: {self.month.currentIndex()}")
        print(f"year index: {self.year.currentIndex()}")

        if (
            self.box1.currentIndex() != 0 and
            self.name.text().strip() and
            self.surname.text().strip() and
            self.section.text().strip() and
            self.code.text().strip() and
            self.day.currentIndex() != 0 and
            self.month.currentIndex() != 0 and
            self.year.currentIndex() != 0
        ):
            self.lbl.setText(f"Note: {self.score}")
            self.lbl1.setText(f"Temps pris: {self.time_taken} secondes")
            self.right.setText(f"Réponses correctes: {self.right_count}")
            self.bhim.setText(f"Mauvaises réponses: {self.wrong_count}")

            self.lbl.show()
            self.lbl1.show()
            self.right.show()
            self.bhim.show()
        else:
            self.lbl.hide()
            self.lbl1.hide()
            self.right.hide()
            self.bhim.hide()

    def save_to_docx(self):
        """ Save the displayed data to a Word (.docx) file with enhanced formatting. """
        name = self.name.text().strip()
        surname = self.surname.text().strip()
        section = self.section.text().strip()
        code = self.code.text().strip()  
        box_value = self.box1.currentText()
        selected_day = self.day.currentText()
        selected_month = self.month.currentText()
        selected_year = self.year.currentText()

        doc = Document()

        title = doc.add_heading("Liste de Note", level=1)
        title.alignment = 1
        title_run = title.add_run()
        title_run.font.size = Pt(18)

        doc.add_paragraph().add_run("-" * 50).font.color.rgb = RGBColor(128, 128, 128)

        def add_styled_paragraph(label, value, font_size=18):
            paragraph = doc.add_paragraph()
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            value_run = paragraph.add_run(value)
            value_run.font.color.rgb = RGBColor(0, 102, 204)
            value_run.font.size = Pt(font_size)

            for run in paragraph.runs:
                run.font.size = Pt(font_size)

            return paragraph

        add_styled_paragraph("Nom", name, 18)
        add_styled_paragraph("Prenom", surname, 18)
        add_styled_paragraph("Section", section, 18)
        add_styled_paragraph("Module code", code, 18)  
        add_styled_paragraph("Examen", box_value, 18)
        add_styled_paragraph("Date", f"{selected_day}/{selected_month}/{selected_year}", 18)
        add_styled_paragraph("Note", str(self.score), 18)
        add_styled_paragraph("Temps pris", f"{self.time_taken} secondes", 18)
        add_styled_paragraph("Réponses correctes", str(self.right_count), 18)
        add_styled_paragraph("Mauvaises réponses", str(self.wrong_count), 18)

        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, 
            "Save File", 
            "", 
            "Word Documents (*.docx)", 
            options=options
        )

        if file_path:
            doc.save(file_path)
            QtWidgets.QMessageBox.information(self, "Success", f"Data saved to {file_path}")
        else:
            QtWidgets.QMessageBox.warning(self, "Cancelled", "Save operation was cancelled.")

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = LastApp(85, 120, 8, 2)  
    window.show()
    sys.exit(app.exec_())
